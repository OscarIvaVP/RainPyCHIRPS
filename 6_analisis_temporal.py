import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches

# Cargar el archivo
file_path = 'output/datos-tabulares/promedio_pp_area.xlsx'
data = pd.ExcelFile(file_path).parse('Sheet1')

# Crear un documento Word
doc = Document()
doc.add_heading('Informe de Análisis Temporal de Precipitación', level=1)

# Texto introductorio
doc.add_paragraph(
    "Este informe presenta un análisis detallado de las precipitaciones mensuales y anuales en diferentes regiones. "
    "Incluye visualizaciones gráficas, análisis de tendencias y categorización de los años según su nivel de precipitación."
)

# Definir etiquetas de meses globalmente
months_labels = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                 'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']

# Función auxiliar para guardar gráficos e insertarlos en el documento Word
def save_plot_and_add_to_doc(title, description, filename, plot_function):
    """Helper function to save plots and add them to Word."""
    plot_function()
    plt.savefig(filename, dpi=300)
    plt.close()
    doc.add_heading(title, level=2)
    doc.add_paragraph(description)
    doc.add_picture(filename, width=Inches(6))

# Gráfico 1: Histórico
def plot_historico():
    data['Fecha'] = pd.to_datetime(data['Fecha'])
    plt.figure(figsize=(15, 10))
    for column in data.columns[1:]:
        plt.plot(data['Fecha'], data[column], label=column)
    plt.xlabel('')
    plt.ylabel('Precipitación mensual (mm)', fontweight='bold')
    plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.05), ncol=7, fontsize='small')
    plt.grid()
    plt.tight_layout()

save_plot_and_add_to_doc(
    "Gráfico Histórico de Precipitación",
    "Este gráfico muestra las precipitaciones mensuales históricas para cada región analizada.",
    "output/graficos/historico.png",
    plot_historico
)

# Gráfico 2: Promedios Mensuales
def plot_promedios_mensuales():
    data['Month'] = data['Fecha'].dt.month
    monthly_averages = data.groupby('Month').mean(numeric_only=True)
    plt.figure(figsize=(15, 10))
    for column in monthly_averages.columns:
        plt.plot(monthly_averages.index, monthly_averages[column], label=column)
    plt.xticks(range(1, 13), months_labels, rotation=45)
    plt.xlabel('')
    plt.ylabel('Promedio de precipitación mensual (mm)', fontweight='bold')
    plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=7, fontsize='small')
    plt.grid()
    plt.tight_layout()

save_plot_and_add_to_doc(
    "Promedios Mensuales de Precipitación",
    "El gráfico muestra los promedios mensuales de precipitación para cada región, agrupados por mes.",
    "output/graficos/promedios_mensuales.png",
    plot_promedios_mensuales
)

# Gráfico 3: Correlaciones
def plot_correlation():
    correlation_matrix = data.drop(columns=['Fecha', 'Month']).corr()
    plt.figure(figsize=(15, 10))
    plt.imshow(correlation_matrix, cmap='coolwarm', aspect='auto')
    plt.colorbar(label='Correlation Coefficient')
    plt.title('Mapa de Calor de Precipitación entre las Áreas', fontweight='bold')
    plt.xticks(range(len(correlation_matrix.columns)), correlation_matrix.columns, rotation=90)
    plt.yticks(range(len(correlation_matrix.index)), correlation_matrix.index)
    plt.tight_layout()

save_plot_and_add_to_doc(
    "Mapa de Correlaciones entre Regiones",
    "Este gráfico muestra las correlaciones entre las precipitaciones de diferentes regiones.",
    "output/graficos/correlaciones.png",
    plot_correlation
)

# Gráfico 4: Totales Anuales
def plot_totales_anuales():
    data['Year'] = data['Fecha'].dt.year
    annual_totals_corrected = data.drop(columns=['Month']).groupby('Year').sum(numeric_only=True)
    plt.figure(figsize=(15, 10))
    for column in annual_totals_corrected.columns:
        plt.plot(annual_totals_corrected.index, annual_totals_corrected[column], label=column)
    plt.xlabel('')
    plt.ylabel('Precipitación anual (mm)', fontweight='bold')
    plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.05), ncol=7, fontsize='small')
    plt.grid()
    plt.tight_layout()

save_plot_and_add_to_doc(
    "Totales Anuales de Precipitación",
    "El gráfico muestra los totales anuales de precipitación para cada región.",
    "output/graficos/totales_anuales.png",
    plot_totales_anuales
)

# Gráfico 5: Boxplots
def plot_boxplots():
    regions = data.columns[1:-2]  # Excluyendo 'Fecha', 'Month', y 'Year'
    monthly_data = [data[region].groupby(data['Month']).apply(list) for region in regions]
    plt.figure(figsize=(18, 12))
    for idx, region in enumerate(regions, 1):
        plt.subplot(len(regions) // 7 + 1, 7, idx)
        plt.boxplot(monthly_data[idx - 1], labels=months_labels)
        plt.title(region, fontsize=10, pad=5, fontweight='bold')
        plt.xticks(rotation=90)
        plt.xlabel('')
        plt.ylabel('mm', fontsize=8, fontweight='bold')
        plt.grid(True)
    plt.suptitle('Boxplots Mensuales de Precipitación por Región', fontsize=16, fontweight='bold')
    plt.tight_layout(rect=[0, 0.03, 1, 0.95])

save_plot_and_add_to_doc(
    "Boxplots Mensuales por Región",
    "Este gráfico muestra las distribuciones mensuales de precipitación para cada región.",
    "output/graficos/boxplots.png",
    plot_boxplots
)

# Gráfico 6: Subplots de Líneas
def plot_subplots_lineas():
    regions = data.columns[1:-2]
    plt.figure(figsize=(18, 12))
    for idx, region in enumerate(regions, 1):
        plt.subplot(len(regions) // 7 + 1, 7, idx)
        monthly_data = data.groupby(data['Fecha'].dt.month)[region].mean()
        plt.plot(monthly_data.index, monthly_data.values, marker='o', label=region)
        plt.title(region, fontsize=10, pad=5, fontweight='bold')
        plt.xticks(range(1, 13), months_labels, rotation=90, fontsize=8)
        plt.yticks(fontsize=8)
        plt.xlabel('')
        plt.ylabel('mm', fontsize=8, fontweight='bold')
        plt.grid(True)
    plt.tight_layout(rect=[0, 0.03, 1, 0.95])

save_plot_and_add_to_doc(
    "Subplots de Líneas Mensuales",
    "Este gráfico muestra las tendencias mensuales de precipitación para cada región como líneas individuales.",
    "output/graficos/lineas_subplots.png",
    plot_subplots_lineas
)

# Gráfico 7: Categorías de Años
def plot_categorias():
    datos = data.drop(['Month', 'Year'], axis=1)
    columnas = [col for col in datos.columns if col != 'Fecha']
    promedio_diario = datos[columnas].mean(axis=1)
    datos['Precipitacion'] = promedio_diario
    datos['Anio'] = datos['Fecha'].dt.year
    precipitacion_anual = datos.groupby('Anio')['Precipitacion'].sum()
    q1 = precipitacion_anual.quantile(0.25)
    q3 = precipitacion_anual.quantile(0.75)
    colors = ['red' if p < q1 else 'green' if p > q3 else 'blue' for p in precipitacion_anual]
    plt.figure(figsize=(18, 12))
    plt.bar(precipitacion_anual.index, precipitacion_anual.values, color=colors)
    plt.xlabel("")
    plt.ylabel("Precipitación anual (mm)", fontweight='bold')
    leyenda_normal = mpatches.Patch(color='blue', label='Año normal')
    leyenda_humedo = mpatches.Patch(color='green', label='Año húmedo')
    leyenda_seco = mpatches.Patch(color='red', label='Año seco')
    plt.legend(handles=[leyenda_seco, leyenda_humedo, leyenda_normal])
    plt.xticks(rotation='vertical')
    plt.grid(True)

save_plot_and_add_to_doc(
    "Categorización de Años Hidrológicos",
    "Este gráfico clasifica los años como secos, normales o húmedos en función de los totales anuales de precipitación.",
    "output/graficos/categorias_anos.png",
    plot_categorias
)

# Guardar el informe
doc.save("output/informes/informe_precipitacion_temporal.docx")
print("Informe generado: 'output//informes/informe_precipitacion_completo.docx'")