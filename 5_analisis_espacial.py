#Librerias
#!pip install python-docx

import pandas as pd
import geopandas as gpd
import matplotlib.pyplot as plt
from matplotlib.colors import Normalize
from matplotlib import cm
import numpy as np
from docx import Document
from docx.shared import Inches

# Crear el documento Word
doc = Document()
doc.add_heading("Informe de Análisis de Precipitación", level=1)

# Ruta del shapefile y Excel
shapefile_path = 'input/area-estudio/area-estudio.shp'
excel_path = 'output/datos-tabulares/promedio_pp_area.xlsx'

###############################################################################
######     GENERACIÓN DE GRÁFICOS POR SEPARADO
###############################################################################

# Gráfico 1: Promedio Anual
gdf_1 = gpd.read_file(shapefile_path)
df_1 = pd.read_excel(excel_path)
df_1['Año'] = pd.to_datetime(df_1['Fecha']).dt.year
numeric_cols_1 = df_1.select_dtypes(include=['float64', 'int64']).columns
annual_totals = df_1.groupby('Año')[numeric_cols_1].sum()
area_avg_totals = annual_totals.mean().reset_index()
area_avg_totals.columns = ['nombre', 'avg_precipitation']
gdf_1 = gdf_1.merge(area_avg_totals, on='nombre', how='left')
fig, ax = plt.subplots(1, 1, figsize=(12, 8), dpi=300)
vmin = gdf_1['avg_precipitation'].min()
vmax = gdf_1['avg_precipitation'].max()
norm = Normalize(vmin=vmin, vmax=vmax)
cmap = cm.Blues
gdf_1.plot(column='avg_precipitation', cmap=cmap, linewidth=0.8, ax=ax, edgecolor='black')
sm = cm.ScalarMappable(cmap=cmap, norm=norm)
sm._A = []
cbar = fig.colorbar(sm, ax=ax, orientation='vertical', fraction=0.03, pad=0.04, shrink=0.7)
cbar.ax.set_title('Promedio histórico\nanual', pad=20, fontsize=10, loc='center')  # Título a lo largo de la barra
cbar.ax.text(0.5, 1.03, 'mm', transform=cbar.ax.transAxes, fontsize=10, ha='center')  # Unidad encima de la barra
ax.axis('off')
plt.tight_layout()
graph1_path = "output/graficos/promedio_anual_historico_espacial.png"
plt.savefig(graph1_path, bbox_inches='tight')  # Guarda ajustando bordes al contenido
plt.close()

# Gráfico 2: Comparación por Años
gdf_3 = gpd.read_file(shapefile_path)
df_3 = pd.read_excel(excel_path)
df_3['Año'] = pd.to_datetime(df_3['Fecha']).dt.year
numeric_cols_3 = df_3.select_dtypes(include=['float64', 'int64']).columns
years = [1981, 2002, 2024]
yearly_totals = {year: df_3[df_3['Año'] == year][numeric_cols_3].sum().reset_index() for year in years}

fig, axes = plt.subplots(1, 3, figsize=(18, 6), dpi=300)
vmin = min(yearly_totals[year][0].min() for year in years)
vmax = max(yearly_totals[year][0].max() for year in years)
norm = Normalize(vmin=vmin, vmax=vmax)

for ax, year in zip(axes, years):
    year_data = yearly_totals[year]
    year_data.columns = ['nombre', 'total_precipitation']
    gdf_year = gdf_3.merge(year_data, on='nombre', how='left')
    gdf_year.plot(column='total_precipitation', cmap=cmap, linewidth=0.8, ax=ax, edgecolor='black')
    ax.set_title(f'{year}', fontsize=14, fontweight='bold')
    ax.axis('off')

# Ajustar márgenes y agregar la barra de colores
fig.subplots_adjust(left=0.05, right=0.95, top=0.9, bottom=0.1, wspace=0.2)
cbar_ax = fig.add_axes([0.2, 0.02, 0.6, 0.04])  # Ajuste de posición de la barra de colores
sm = cm.ScalarMappable(cmap=cmap, norm=norm)
sm._A = []

# Configurar la barra de colores
cbar = fig.colorbar(sm, cax=cbar_ax, orientation='horizontal')
tick_values = np.linspace(vmin, vmax, 7)  # Generar 7 valores equiespaciados
cbar.set_ticks(tick_values)
cbar.ax.set_xticklabels([f'{tick:.2f}' for tick in tick_values], fontsize=12)  # Aumentar tamaño de fuente de etiquetas
cbar.set_label('Total Anual de Precipitación (mm)', fontsize=14)  # Etiqueta más grande

plt.tight_layout()  # Optimiza el espacio en el gráfico
graph2_path = "output/graficos/comparación_anual_espacial.png"
plt.savefig(graph2_path, bbox_inches='tight')  # Guarda ajustando bordes al contenido
plt.close()

# Gráfico 3: Promedio Mensual

gdf_2 = gpd.read_file(shapefile_path)
df_2 = pd.read_excel(excel_path)

df_2['Mes'] = pd.to_datetime(df_2['Fecha']).dt.month

numeric_cols_2 = df_2.select_dtypes(include=['float64', 'int64']).columns

monthly_avg = df_2.groupby('Mes')[numeric_cols_2].mean()

fig, axes = plt.subplots(4, 3, figsize=(15, 10), dpi=300)
axes = axes.flatten()
vmin = monthly_avg.min().min()
vmax = monthly_avg.max().max()
norm = Normalize(vmin=vmin, vmax=vmax)

for month, ax in enumerate(axes, start=1):
    if month in monthly_avg.index:
        month_data = monthly_avg.loc[month].reset_index()
        month_data.columns = ['nombre', 'avg_precipitation']
        gdf_month = gdf_2.merge(month_data, on='nombre', how='left')
        gdf_month.plot(column='avg_precipitation', cmap=cmap, linewidth=0.5, ax=ax, edgecolor='black')
        ax.set_title(f'Mes {month}', fontsize=14, fontweight='bold')
    ax.axis('off')

fig.subplots_adjust(left=0.05, right=0.95, top=0.9, bottom=0.1, hspace=0.3, wspace=0.2)
cbar_ax = fig.add_axes([0.25, -0.05, 0.5, 0.02])  # Ajuste preciso de la barra de colores
sm = cm.ScalarMappable(cmap=cmap, norm=norm)
sm._A = []
cbar = fig.colorbar(sm, cax=cbar_ax, orientation='horizontal')
tick_values = np.linspace(vmin, vmax, 5)
cbar.set_ticks(tick_values)
cbar.ax.set_xticklabels([f'{tick:.2f}' for tick in tick_values])
cbar.set_label('Promedio Mensual de Precipitación (mm)', fontsize=14)
plt.tight_layout()  # Optimiza el espacio en el gráfico
graph3_path = "output/graficos/promedio_mensual_historio_espacial.png"
plt.savefig(graph3_path, bbox_inches='tight')  # Guarda ajustando bordes al contenido
plt.close()

###############################################################################
######     AGREGAR GRÁFICOS AL DOCUMENTO WORD
###############################################################################

# Agregar Gráfico 1
doc.add_heading("Gráfico 1: Promedio Anual de Totales de Precipitación", level=2)
doc.add_paragraph(
    "Este gráfico muestra el promedio anual de totales de precipitación en el área de estudio."
)
doc.add_picture(graph1_path, width=Inches(6))

# Agregar Gráfico 2
doc.add_heading("Gráfico 2: Comparación de Totales de Precipitación", level=2)
doc.add_paragraph(
    "Este gráfico muestra los totales de precipitación anual para los años 1981, 2000 y 2023."
)
doc.add_picture(graph2_path, width=Inches(6))

# Agregar Gráfico 3
doc.add_heading("Gráfico 3: Promedio Mensual de Precipitación", level=2)
doc.add_paragraph(
    "Este gráfico muestra los promedios mensuales de precipitación para el área de estudio."
)
doc.add_picture(graph3_path, width=Inches(6))

###############################################################################
######     GUARDAR EL INFORME
###############################################################################

report_path = "output/informes/informe_precipitacion_espacial.docx"
doc.save(report_path)
print(f"Informe generado: {report_path}")




