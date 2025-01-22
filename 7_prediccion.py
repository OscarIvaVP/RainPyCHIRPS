import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import GridSearchCV, TimeSeriesSplit
from sklearn.preprocessing import MinMaxScaler
from scipy.stats import norm
from docx import Document
from docx.shared import Inches
import os
import random

# Función para calcular intervalos de confianza
def confidence_intervals(predictions, std_dev, confidence=0.95):
    z = norm.ppf((1 + confidence) / 2)
    lower_bound = np.maximum(predictions - z * std_dev, 0)  # Limita el valor mínimo a 0
    upper_bound = predictions + z * std_dev
    return lower_bound, upper_bound

# Crear carpeta de salida si no existe
output_dir = "output"
os.makedirs(output_dir, exist_ok=True)

# Cargar los datos
data = pd.read_excel(f'{output_dir}/datos-tabulares/promedio_pp_area.xlsx')
data['Fecha'] = pd.to_datetime(data['Fecha'])
data.set_index('Fecha', inplace=True)
data_cleaned = data.dropna()

# Agregar características adicionales
data_cleaned['Mes'] = data_cleaned.index.month
data_cleaned['Año'] = data_cleaned.index.year
data_cleaned['Seno'] = np.sin(2 * np.pi * data_cleaned.index.month / 12)
data_cleaned['Coseno'] = np.cos(2 * np.pi * data_cleaned.index.month / 12)

# Variables
window_size = 12
results = []
metrics = []
future_results = {}

# Modelado por cada región
for column_name in data_cleaned.columns[:-4]:
    print(f"Procesando región: {column_name}")

    values = data_cleaned[column_name].values.reshape(-1, 1)
    dates = data_cleaned.index

    # Escalado
    scaler = MinMaxScaler(feature_range=(0, 1))
    scaled_values = scaler.fit_transform(values)

    # Construcción de ventanas
    X, y = [], []
    for i in range(len(scaled_values) - window_size):
        window = scaled_values[i:i + window_size].flatten()
        month_features = data_cleaned.iloc[i + window_size][['Mes', 'Año', 'Seno', 'Coseno']].values
        X.append(np.concatenate([window, month_features]))
        y.append(scaled_values[i + window_size])
    X, y = np.array(X), np.array(y).flatten()

    # División de datos en entrenamiento y validación
    train_size = int(len(X) * 0.8)
    X_train, X_test = X[:train_size], X[train_size:]
    y_train, y_test = y[:train_size], y[train_size:]
    train_dates = dates[window_size:train_size + window_size]
    test_dates = dates[train_size + window_size:]

    # Grid Search para encontrar mejores hiperparámetros
    param_grid = {
        'n_estimators': [100, 200],
        'max_depth': [10, 20, None],
        'min_samples_split': [2, 5],
        'min_samples_leaf': [1, 2]
    }
    tscv = TimeSeriesSplit(n_splits=3)
    grid_search = GridSearchCV(RandomForestRegressor(random_state=42), param_grid, cv=tscv, scoring='neg_mean_squared_error', n_jobs=-1)
    grid_search.fit(X_train, y_train)
    model = grid_search.best_estimator_

    # Predicciones y métricas
    y_train_pred = model.predict(X_train)
    y_test_pred = model.predict(X_test)

    y_train_inverse = scaler.inverse_transform(y_train.reshape(-1, 1)).flatten()
    y_train_pred_inverse = scaler.inverse_transform(y_train_pred.reshape(-1, 1)).flatten()
    y_test_inverse = scaler.inverse_transform(y_test.reshape(-1, 1)).flatten()
    y_test_pred_inverse = scaler.inverse_transform(y_test_pred.reshape(-1, 1)).flatten()

    # Resultados
    results.extend([
        {"Fecha": date, "Valor": value, "Simulado": pred, "Conjunto": "Entrenamiento", "Región": column_name}
        for date, value, pred in zip(train_dates, y_train_inverse, y_train_pred_inverse)
    ])
    results.extend([
        {"Fecha": date, "Valor": value, "Simulado": pred, "Conjunto": "Validación", "Región": column_name}
        for date, value, pred in zip(test_dates, y_test_inverse, y_test_pred_inverse)
    ])

    metrics.append({
        "Región": column_name,
        "MSE_Entrenamiento": round(mean_squared_error(y_train_inverse, y_train_pred_inverse), 2),
        "MAE_Entrenamiento": round(mean_absolute_error(y_train_inverse, y_train_pred_inverse), 2),
        "R2_Entrenamiento": round(r2_score(y_train_inverse, y_train_pred_inverse), 2),
        "MSE_Validación": round(mean_squared_error(y_test_inverse, y_test_pred_inverse), 2),
        "MAE_Validación": round(mean_absolute_error(y_test_inverse, y_test_pred_inverse), 2),
        "R2_Validación": round(r2_score(y_test_inverse, y_test_pred_inverse), 2)
    })

    # Predicciones futuras hasta 2030
    future_dates = pd.date_range(start=data_cleaned.index[-1], end='2030-12-31', freq='M')
    last_window = scaled_values[-window_size:].flatten()
    future_predictions = []
    lower_bounds = []
    upper_bounds = []

    for future_date in future_dates:
        future_features = [future_date.month, future_date.year, np.sin(2 * np.pi * future_date.month / 12), np.cos(2 * np.pi * future_date.month / 12)]
        pred_input = np.concatenate([last_window, future_features])
        pred = model.predict(pred_input.reshape(1, -1))[0]
        lower_bound, upper_bound = confidence_intervals(pred, std_dev=0.1 * pred)
        random_pred = random.uniform(lower_bound, upper_bound)
        future_predictions.append(scaler.inverse_transform([[random_pred]])[0, 0])
        lower_bounds.append(scaler.inverse_transform([[lower_bound]])[0, 0])
        upper_bounds.append(scaler.inverse_transform([[upper_bound]])[0, 0])
        last_window = np.roll(last_window, -1)
        last_window[-1] = pred

    future_results[column_name] = {
        "Fechas": future_dates,
        "Predicciones": future_predictions,
        "Límite Inferior": lower_bounds,
        "Límite Superior": upper_bounds
    }

results_df = pd.DataFrame(results)
metrics_df = pd.DataFrame(metrics)

# Guardar métricas en Excel
with pd.ExcelWriter(f'{output_dir}/datos-tabulares/analisis_random_forest.xlsx') as writer:
    metrics_df.to_excel(writer, sheet_name="Métricas", index=False)
    results_df.to_excel(writer, sheet_name="Observados y Simulados", index=False)

# Graficar resultados
num_regions = len(data_cleaned.columns[:-4])
cols = 3
rows = (num_regions + cols - 1) // cols

fig_train, axes_train = plt.subplots(rows, cols, figsize=(15, rows * 4), sharex=False, sharey=False)
fig_valid, axes_valid = plt.subplots(rows, cols, figsize=(15, rows * 4), sharex=False, sharey=False)

axes_train = axes_train.flatten()
axes_valid = axes_valid.flatten()

for idx, column_name in enumerate(data_cleaned.columns[:-4]):
    region_data = results_df[results_df["Región"] == column_name]
    train_data = region_data[region_data["Conjunto"] == "Entrenamiento"]
    validation_data = region_data[region_data["Conjunto"] == "Validación"]

    ax_train = axes_train[idx]
    ax_train.plot(train_data["Fecha"], train_data["Valor"], label="Observado", color="black")
    ax_train.plot(train_data["Fecha"], train_data["Simulado"], label="Simulado", color="red", linestyle="--")
    ax_train.set_title(f"Entrenamiento: {column_name}", fontsize=10)
    ax_train.legend(loc="upper right")  # Agregar leyenda
    ax_train.grid(True)

    ax_valid = axes_valid[idx]
    ax_valid.plot(validation_data["Fecha"], validation_data["Valor"], label="Observado", color="black")
    ax_valid.plot(validation_data["Fecha"], validation_data["Simulado"], label="Simulado", color="red", linestyle="--")
    ax_valid.plot(future_results[column_name]["Fechas"], future_results[column_name]["Predicciones"], label="Predicción", color="blue")
    ax_valid.fill_between(
        future_results[column_name]["Fechas"],
        future_results[column_name]["Límite Inferior"],
        future_results[column_name]["Límite Superior"],
        color="blue",
        alpha=0.2,
        label="Intervalo de confianza"
    )
    ax_valid.set_title(f"Validación y Predicción: {column_name}", fontsize=10)
    ax_valid.legend(loc="upper right")  # Agregar leyenda
    ax_valid.grid(True)

# Eliminar subplots vacíos
for idx in range(len(data_cleaned.columns[:-4]), len(axes_train)):
    fig_train.delaxes(axes_train[idx])
for idx in range(len(data_cleaned.columns[:-4]), len(axes_valid)):
    fig_valid.delaxes(axes_valid[idx])

fig_train.tight_layout()
fig_valid.tight_layout()

fig_train.savefig(f"{output_dir}/graficos/train_plot.png")
fig_valid.savefig(f"{output_dir}/graficos/valid_plot.png")

# Crear documento Word
doc = Document()
doc.add_heading("Informe de Resultados: Predicción de Precipitación", level=1)
doc.add_paragraph("Este informe presenta los resultados obtenidos.")

# Insertar tabla de métricas al final
doc.add_heading("Métricas de Evaluación", level=2)
doc.add_paragraph("La siguiente tabla muestra las métricas de evaluación de los modelos.")
table = doc.add_table(rows=1, cols=len(metrics_df.columns))
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(metrics_df.columns):
    hdr_cells[i].text = col_name

for _, row in metrics_df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)
        
# Insertar gráficas en el informe
doc.add_heading("Gráficas de Entrenamiento", level=2)
doc.add_picture(f"{output_dir}/graficos/train_plot.png", width=Inches(6))

doc.add_heading("Gráficas de Validación y Predicción", level=2)
doc.add_picture(f"{output_dir}/graficos/valid_plot.png", width=Inches(6))

# Guardar el informe en Word
doc.save(f"{output_dir}/informes/informe_prediccion.docx")
print("Informe generado exitosamente: 'output/informes/informe_prediccion.docx'")







