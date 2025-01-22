# -*- coding: utf-8 -*-

# Importar librerías necesarias
# pip install geopandas rasterio python-docx

import os
import geopandas as gpd
import rasterio
from rasterio.mask import mask
from shapely.geometry import box
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from rasterio.plot import show
from docx import Document
from docx.shared import Inches

# Crear carpeta de salida si no existe
output_folder = "output"
os.makedirs(output_folder, exist_ok=True)

# Cargar shapefile
shapefile_path = "input/area-estudio/area-estudio.shp"
assert os.path.exists(shapefile_path), f"Archivo no encontrado: {shapefile_path}"
gdf = gpd.read_file(shapefile_path)

# Generar primer gráfico
plt.figure()
gdf.plot()
plt.title('Shapefile: área de estudio')
plt.savefig(f"{output_folder}/graficos/area-estudio.png", dpi=300, bbox_inches='tight')
plt.close()

# Crear DataFrame limpio para almacenar resultados
results_df = pd.DataFrame(columns=['nombre', 'Fecha', 'Promedio_Precipitacion'])

# Definir carpeta de rasters
raster_folder = "output/imagenes-chirps-mes"
assert os.path.exists(raster_folder), f"Carpeta no encontrada: {raster_folder}"

# Distancia del buffer (ajustar según el CRS, por ejemplo, 0.01 para grados o 1000 para metros)
buffer_distance = 0.01  # Aprox. 1 km si CRS está en grados (WGS 84)

# Función para extraer fecha del nombre del archivo raster
def extract_date_from_filename(filename):
    # Asume que el nombre tiene formato "chirps-v2.0.AAAA.MM.tif"
    parts = filename.split('.')
    year = int(parts[-3])  # AAAA
    month = int(parts[-2])  # MM
    return pd.Timestamp(year=year, month=month, day=1)

# Iterar sobre cada archivo raster en la carpeta
for raster_file in os.listdir(raster_folder):
    if raster_file.endswith('.tif'):
        raster_path = os.path.join(raster_folder, raster_file)
        
        with rasterio.open(raster_path) as src:
            # Extraer la fecha del nombre del archivo
            raster_date = extract_date_from_filename(raster_file)
            
            # Asegurar que CRS del shapefile coincida con el raster
            if gdf.crs != src.crs:
                gdf = gdf.to_crs(src.crs)
            
            # Iterar sobre cada polígono en el shapefile
            for index, row in gdf.iterrows():
                municipality_name = row['nombre']
                polygon_shape = row['geometry']

                # Verificar si el polígono está vacío
                if polygon_shape.is_empty:
                    print(f"El polígono {municipality_name} está vacío.")
                    continue
                
                # Ampliar el polígono con un buffer
                polygon_shape_buffered = polygon_shape.buffer(buffer_distance)

                # Verificar si el polígono ampliado intersecta con el raster
                raster_bounds = box(*src.bounds)
                if not polygon_shape_buffered.intersects(raster_bounds):
                    print(f"El polígono {municipality_name} no intersecta con el raster {raster_file}.")
                    continue
                
                try:
                    # Aplicar máscara con el polígono ampliado
                    out_image, _ = mask(src, [polygon_shape_buffered], crop=True)
                    
                    # Calcular promedio de precipitación, ignorando valores no válidos (<= 0)
                    avg_precipitation = np.nanmean(out_image[out_image > 0])
                    
                    # Si el promedio es NaN, indicar que no hubo datos válidos
                    if np.isnan(avg_precipitation):
                        print(f"Sin datos válidos para {municipality_name} en {raster_file}.")
                        continue
                    
                    # Añadir resultado al DataFrame
                    results_df = pd.concat([results_df, pd.DataFrame([{
                        'nombre': municipality_name,
                        'Fecha': raster_date,
                        'Promedio_Precipitacion': avg_precipitation
                    }])], ignore_index=True)
                except ValueError as e:
                    print(f"Error al procesar {municipality_name} en {raster_file}: {str(e)}")

# Reorganizar el DataFrame para que las columnas sean las subcuencas
pivot_df = results_df.pivot(index='Fecha', columns='nombre', values='Promedio_Precipitacion')

# Guardar resultados en un archivo Excel
excel_path = os.path.join(output_folder, 'datos-tabulares/promedio_pp_area.xlsx')
pivot_df.to_excel(excel_path)
print("Guardado con éxito.")

# Visualización de un raster específico con el shapefile
specific_raster_path = os.path.join(raster_folder, "chirps-v2.0.1981.01.tif")
assert os.path.exists(specific_raster_path), f"Raster no encontrado: {specific_raster_path}"

with rasterio.open(specific_raster_path) as src:
    fig, ax = plt.subplots(dpi=300)
    
    # Mostrar el raster
    show(src, ax=ax, cmap='gray')
    
    # Mostrar los límites originales del shapefile
    gdf.boundary.plot(ax=ax, color='red', label='Original')
    
    # Mostrar los polígonos ampliados con buffer
    gdf.geometry.apply(lambda x: x.buffer(buffer_distance)).plot(ax=ax, color='blue', alpha=0.3, label='Buffer')
    
    # Añadir leyendas y títulos
    plt.legend()
    plt.title('Lugar del Shapefile y el mundo')
    plt.xlabel('Longitud')
    plt.ylabel('Latitud')
    plt.savefig(f"{output_folder}/graficos/ubicacion_area-estudio.png", dpi=300, bbox_inches='tight')
    plt.close()

# Crear el informe en Word
word_path = os.path.join(output_folder, "informes/informe_area_estudio.docx")
doc = Document()

# Título del documento
doc.add_heading("Informe de Ubicación de Análisis", level=1)

# Introducción y primer gráfico
doc.add_heading("1. Shapefile: Área de Estudio", level=2)
doc.add_paragraph(
    "Este gráfico muestra el área de estudio utilizada en el análisis. Los polígonos corresponden a las áreas seleccionadas para evaluar la precipitación promedio.")
doc.add_picture(f"{output_folder}/graficos/area-estudio.png", width=Inches(5))

# Segundo gráfico y descripción
doc.add_heading("2. Ubicación del Shapefile en el Mundo", level=2)
doc.add_paragraph(
    "Este gráfico muestra la ubicación del shapefile en el contexto global, junto con una visualización del raster de precipitación utilizado para el análisis.")
doc.add_picture(f"{output_folder}/graficos/ubicacion_area-estudio.png", width=Inches(5))

# Guardar el informe
doc.save(word_path)
print(f"Informe guardado en: {word_path}")